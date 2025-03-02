import streamlit as st
import os
import pandas as pd
from datetime import datetime, timedelta
from docx import Document
from db_config2 import get_db_connection, insert_default_attendance
from take_attendace import convert_24_to_12
import pymysql

# Constants
KNOWN_FACES_DIR = "known_faces"
ALLOWED_EXTENSIONS = {'.jpg', '.jpeg', '.png'}
USERNAME = "admin"
PASSWORD = "admin"

# Ensure the directory exists
os.makedirs(KNOWN_FACES_DIR, exist_ok=True)

def handle_user_login():
    """Handle login and ensure attendance is initialized"""
    insert_default_attendance()

def login():
    """Login page for authentication"""
    st.header("🔒 Login")
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")
    if st.button("Login"):
        if username == USERNAME and password == PASSWORD:
            st.session_state.logged_in = True
            st.success("Login successful!")
            handle_user_login()
        else:
            st.error("Invalid username or password")

def register_student():
    """Student registration page"""
    st.header("📸 Register New Student")
    
    student_name = st.text_input("Enter Student's Name:", help="Enter the full name of the student")
    uploaded_image = st.file_uploader("Upload Image", type=list(ext[1:] for ext in ALLOWED_EXTENSIONS),
                                      help="Upload a clear frontal face photo")
    
    if st.button("Register", help="Click to register the student"):
        if not student_name.strip():
            st.error("Please enter a student name.")
            return
        if not uploaded_image:
            st.error("Please upload an image.")
            return
            
        try:
            # Save the image to the known_faces folder
            filepath = os.path.join(KNOWN_FACES_DIR, f"{student_name}.jpg")
            with open(filepath, "wb") as f:
                uploaded_image.seek(0)
                f.write(uploaded_image.read())

            st.success(f"Student '{student_name}' registered successfully!")
            st.image(filepath, caption=f"Registered image for {student_name}", use_column_width=True)

            # Insert student into userDetails
            connection = get_db_connection()
            if connection is None:
                st.error("Failed to connect to the database.")
                return

            try:
                with connection.cursor() as cursor:
                    cursor.execute("""
                        INSERT INTO userDetails (name) VALUES (%s)
                        ON DUPLICATE KEY UPDATE name = name;
                    """, (student_name,))

                st.info(f"Student '{student_name}' added to the database.")
            except pymysql.Error as e:
                st.error(f"Failed to insert student into database: {e}")
            finally:
                connection.close()

        except Exception as e:
            st.error(f"Registration failed: {str(e)}")

def view_attendance_data():
    """View and export attendance data"""
    st.header("📊 View Attendance Data")

    # Initialize session state for start_date and end_date
    if 'start_date' not in st.session_state:
        st.session_state.start_date = datetime.now() - timedelta(days=7)
    if 'end_date' not in st.session_state:
        st.session_state.end_date = datetime.now()

    # Add a "Today" button
    if st.button("Today"):
        st.session_state.start_date = datetime.now().date()
        st.session_state.end_date = datetime.now().date()

    # Date range filter
    st.subheader("Filter by Date Range")
    col1, col2 = st.columns(2)
    with col1:
        start_date = st.date_input("Start Date", st.session_state.start_date)
    with col2:
        end_date = st.date_input("End Date", st.session_state.end_date)

    # Update session state with the selected dates
    st.session_state.start_date = start_date
    st.session_state.end_date = end_date

    if start_date > end_date:
        st.error("Start date cannot be after end date")
        return
        
    conn = get_db_connection()
    if conn:
        try:
            with conn.cursor() as cursor:
                query = """
                SELECT 
                    name, date, 
                    entry1_in, entry1_out, entry1_hours,
                    entry2_in, entry2_out, entry2_hours,
                    entry3_in, entry3_out, entry3_hours,
                    entry4_in, entry4_out, entry4_hours,
                    entry5_in, entry5_out, entry5_hours,
                    total_hours, status
                FROM attendance
                WHERE date BETWEEN %s AND %s
                ORDER BY date DESC, name ASC
                """
                cursor.execute(query, (start_date, end_date))
                rows = cursor.fetchall()
                
                if rows:
                    df = pd.DataFrame(rows)
                    
                    # Convert 24-hour format to 12-hour format for display
                    time_columns = [
                        'entry1_in', 'entry1_out', 'entry2_in', 'entry2_out',
                        'entry3_in', 'entry3_out', 'entry4_in', 'entry4_out',
                        'entry5_in', 'entry5_out'
                    ]
                    for col in time_columns:
                        if col in df.columns:
                            df[col] = df[col].apply(lambda x: convert_24_to_12(x) if x else "")
                    
                    selected_status = st.multiselect(
                        "Filter by Status",
                        options=sorted(df['status'].unique()),
                        default=[]
                    )
                    
                    if selected_status:
                        df = df[df['status'].isin(selected_status)]
                    
                    st.dataframe(df, use_container_width=True)
                    
                    st.subheader("Summary Statistics")
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.metric("Total Students", len(df['name'].unique()))
                    with col2:
                        st.metric("Total Days", len(df['date'].unique()))
                    with col3:
                        present_count = len(df[df['status'] == 'Present'])
                        st.metric("Total Present", present_count)
                    with col4:
                        attendance_rate = (present_count / len(df) * 100) if len(df) > 0 else 0
                        st.metric("Attendance Rate", f"{attendance_rate:.1f}%")
                    
                    export_format = st.selectbox("Export Format:", ["Excel", "CSV", "Word"])
                    
                    if st.button("Generate Report"):
                        generate_report(df, export_format, start_date, end_date)
                else:
                    st.info("No records found for the selected date range")
                    
        except Exception as e:
            st.error(f"Error fetching records: {str(e)}")
        finally:
            conn.close()
    else:
        st.error("Database connection failed")

def generate_report(df, export_format, start_date, end_date):
    """Generate attendance report in specified format"""
    try:
        if export_format == "Excel":
            df.to_excel("Attendance_Report.xlsx", index=False)
            with open("Attendance_Report.xlsx", "rb") as f:
                st.download_button("Download Excel Report", data=f, file_name=f"Attendance_Report_{start_date}to{end_date}.xlsx")
        
        elif export_format == "CSV":
            csv = df.to_csv(index=False)
            st.download_button("Download CSV Report", data=csv, file_name=f"Attendance_Report_{start_date}to{end_date}.csv")
        
        else:  # Word format
            doc = Document()
            doc.add_heading(f"Attendance Report ({start_date} to {end_date})", 0)
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(f"Total Students: {len(df['name'].unique())}")
            doc.add_paragraph(f"Total Days: {len(df['date'].unique())}")
            doc.add_paragraph(f"Total Present: {len(df[df['status'] == 'Present'])}")
            doc.save("Attendance_Report.docx")
            with open("Attendance_Report.docx", "rb") as f:
                st.download_button("Download Word Report", data=f, file_name=f"Attendance_Report_{start_date}to{end_date}.docx")
    except Exception as e:
        st.error(f"Error generating report: {str(e)}")

def main():
    """Main function"""
    st.set_page_config(page_title="Manage Students", layout="wide")
    if 'logged_in' not in st.session_state:
        st.session_state.logged_in = False

    if not st.session_state.logged_in:
        login()
    else:
        st.title("Manage Students")
        page = st.sidebar.selectbox("Choose a page:", ["Register Students", "View Attendance Data"])
        if page == "Register Students":
            register_student()
        else:
            view_attendance_data()

if __name__ == "__main__":
    main()
